import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import gspread
from google.oauth2.service_account import Credentials
from datetime import datetime

st.set_page_config(page_title="Asignación de Plazas", page_icon="📋", layout="wide")

st.markdown("""
<style>
    .main-header { background: linear-gradient(135deg, #1a3a5c 0%, #2563a8 100%); color: white; padding: 1.5rem 2rem; border-radius: 10px; margin-bottom: 1.5rem; }
    .main-header h1 { color: white; margin: 0; font-size: 1.6rem; }
    .main-header p { color: #c8d8f0; margin: 0.2rem 0 0 0; font-size: 0.9rem; }
    .metric-card { background: white; border: 1px solid #e2e8f0; border-radius: 10px; padding: 1rem 1.2rem; text-align: center; box-shadow: 0 1px 4px rgba(0,0,0,0.06); }
    .metric-card .value { font-size: 2rem; font-weight: 700; color: #1a3a5c; }
    .metric-card .label { font-size: 0.8rem; color: #64748b; margin-top: 0.2rem; }
    .hora-bar-container { background: #e2e8f0; border-radius: 20px; height: 14px; overflow: hidden; margin-top: 0.3rem; }
    .hora-bar { height: 100%; border-radius: 20px; }
    .plaza-item { background: white; border: 1px solid #e2e8f0; border-radius: 8px; padding: 0.7rem 1rem; margin-bottom: 0.5rem; }
    .plaza-item.asignada { border-left: 4px solid #2563a8; background: #f0f6ff; }
    .plaza-item.nueva { border-left: 4px solid #10b981; }
    .plaza-item.disponible { border-left: 4px solid #94a3b8; }
    .badge { display: inline-block; padding: 0.15rem 0.5rem; border-radius: 20px; font-size: 0.72rem; font-weight: 600; margin-left: 0.4rem; }
    .badge-actual { background: #dbeafe; color: #1e40af; }
    .badge-nueva { background: #d1fae5; color: #065f46; }
    .oficio-preview { background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; padding: 1rem; margin-bottom: 0.5rem; font-size: 0.85rem; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>📋 Sistema de Asignación de Plazas</h1>
    <p>Dirección de Formación Continua — Secretaría de Educación del Estado de Jalisco</p>
</div>
""", unsafe_allow_html=True)

# ─── GOOGLE SHEETS ─────────────────────────────
SPREADSHEET_ID = "1GExm22h2VofySyVqNjdh0509m_zzLiQKVxiXSqlkn_I"
SCOPES = ["https://www.googleapis.com/auth/spreadsheets",
          "https://www.googleapis.com/auth/drive"]

@st.cache_resource
def get_client():
    creds = Credentials.from_service_account_info(
        st.secrets["google_sheets"], scopes=SCOPES)
    return gspread.authorize(creds)

@st.cache_data(ttl=300)
def cargar_sheets():
    gc = get_client()
    sh = gc.open_by_key(SPREADSHEET_ID)
    dv = pd.DataFrame(sh.worksheet("VACANCIA").get_all_records(numericise_ignore=["all"]))
    dc = pd.DataFrame(sh.worksheet("COBERTURA").get_all_records(numericise_ignore=["all"]))
    dh = pd.DataFrame(sh.worksheet("HISTORIAL").get_all_records(numericise_ignore=["all"]))
    return dv, dc, dh

def guardar_asignaciones(oficios, fecha_desde, fecha_hasta):
    try:
        gc = get_client()
        sh = gc.open_by_key(SPREADSHEET_ID)
        try:
            ws = sh.worksheet("ASIGNACIONES")
        except Exception:
            ws = sh.add_worksheet("ASIGNACIONES", rows=1000, cols=10)
            ws.append_row(["FECHA_HORA", "FOLIO", "EMPLEADO", "RFC",
                           "CLAVE_PRESUPUESTAL", "CARGA_HORARIA",
                           "SUSTITUYE_A", "CCT", "EFECTO_INICIAL", "EFECTO_FINAL"])
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        rows = []
        for o in oficios:
            if not o.get("es_nueva"):
                continue
            e = o["emp"]
            rows.append([now, o["folio"], e.get("NOMBRE_INTERINO", ""),
                         e.get("RFC", ""), o["claves_presupuestales"],
                         o["carga_horaria"], o["sustituye_a"],
                         o["clave_cct"], fecha_desde, fecha_hasta])
        if rows:
            ws.append_rows(rows)
        return True
    except Exception as ex:
        st.warning(f"No se pudo guardar en Sheets: {ex}")
        return False

# ─── SIDEBAR ───────────────────────────────────
with st.sidebar:
    st.subheader("📄 Plantilla Word")
    f_plantilla = st.file_uploader("Plantilla de propuesta (.docx)", type=["docx"], key="plt")
    st.divider()
    st.subheader("⚙️ Configuración")
    fecha_desde = st.text_input("Efectos a partir del", value="01 de marzo de 2026")
    fecha_hasta = st.text_input("Hasta", value="31 de mayo de 2026")
    nombre_firmante = st.text_input("Firmante", value="CARMEN YOLANDA QUINTERO REYES")
    cargo_firmante = st.text_input("Cargo", value="SUBSECRETARIA DE FORMACIÓN Y ATENCIÓN AL MAGISTERIO")
    st.divider()
    if st.button("🔄 Recargar datos de Sheets", use_container_width=True):
        st.cache_data.clear()
        st.rerun()

# ─── CARGA DESDE SHEETS ────────────────────────
with st.spinner("Cargando datos desde Google Sheets..."):
    try:
        df_vac, df_cob, df_his = cargar_sheets()
    except Exception as e:
        st.error(f"Error conectando a Google Sheets: {e}")
        st.stop()

# ─── NORMALIZACIÓN ─────────────────────────────
for df in [df_vac, df_cob, df_his]:
    df.columns = df.columns.str.strip().str.upper()

def rc(df, buscar, nuevo):
    for b in buscar:
        m = [c for c in df.columns if b.upper() in c.upper()]
        if m:
            return df.rename(columns={m[0]: nuevo})
    return df

df_vac = rc(df_vac, ["CARGA HORARIA","HORAS"], "CARGA_HORARIA")
df_vac = rc(df_vac, ["CLAVE PRESUPUESTAL","CLAVE PRES"], "CLAVE_PRESUPUESTAL")
df_vac = rc(df_vac, ["SUSTITUYE A","SUSTITUYE"], "SUSTITUYE_A")
df_vac = rc(df_vac, ["NOMBRE CCT","NOMBRE DEL CCT"], "NOMBRE_CCT")
df_vac = rc(df_vac, ["CLAVE CCT","C.C.T."], "CLAVE_CCT")
df_vac = rc(df_vac, ["MOTIVO VACANTE","MOTIVO"], "MOTIVO_VACANTE")
df_vac = rc(df_vac, ["TIPO DE ALTA","TIPO ALTA"], "TIPO_ALTA")

df_cob = rc(df_cob, ["NOMBRE INTERINO","NOMBRE"], "NOMBRE_INTERINO")
df_cob = rc(df_cob, ["CARGA HORARIA","HORAS"], "CARGA_HORARIA")
df_cob = rc(df_cob, ["CLAVE PRESUPUESTAL","CLAVE PRES"], "CLAVE_PRESUPUESTAL")
df_cob = rc(df_cob, ["SUSTITUYE A","SUSTITUYE"], "SUSTITUYE_A")
df_cob = rc(df_cob, ["NOMBRE CCT","NOMBRE DEL CCT"], "NOMBRE_CCT")
df_cob = rc(df_cob, ["CLAVE CCT","C.C.T."], "CLAVE_CCT")
df_cob = rc(df_cob, ["NÚMERO SEGURO","NUMERO SEGURO","NSS"], "NSS")
df_cob = rc(df_cob, ["TELÉFONO","TELEFONO"], "TELEFONO")
df_cob = rc(df_cob, ["CORREO"], "CORREO")
df_cob = rc(df_cob, ["CODIGO POSTAL","CODICO POSTAL","C.P."], "CP")
df_cob = rc(df_cob, ["MOTIVO VACANTE","MOTIVO"], "MOTIVO_VACANTE")

df_his = rc(df_his, ["NOMBRE INTERINO","NOMBRE"], "NOMBRE_INTERINO")
df_his = rc(df_his, ["CARGA DE CUBRIAN","CARGA QUE CUBRIAN","CARGA HORARIA","HORAS"], "HORAS_HIST")
df_his = rc(df_his, ["NÚMERO SEGURO","NUMERO SEGURO","NSS"], "NSS")
df_his = rc(df_his, ["TELÉFONO","TELEFONO"], "TELEFONO")
df_his = rc(df_his, ["CORREO"], "CORREO")
df_his = rc(df_his, ["CODIGO POSTAL","CODICO POSTAL","C.P."], "CP")

df_cob = df_cob.dropna(subset=["NOMBRE_INTERINO"])
df_cob["NOMBRE_INTERINO"] = df_cob["NOMBRE_INTERINO"].astype(str).str.strip().str.upper()
df_his["NOMBRE_INTERINO"] = df_his["NOMBRE_INTERINO"].astype(str).str.strip().str.upper()
df_vac["CLAVE_PRESUPUESTAL"] = df_vac["CLAVE_PRESUPUESTAL"].astype(str).str.strip()

def limpiar_horas(serie):
    return pd.to_numeric(serie.astype(str).str.extract(r'(\d+)', expand=False), errors="coerce").fillna(0).astype(int)

df_vac["CARGA_HORARIA"] = limpiar_horas(df_vac["CARGA_HORARIA"])
df_cob["CARGA_HORARIA"] = limpiar_horas(df_cob["CARGA_HORARIA"])
df_his["HORAS_HIST"]    = limpiar_horas(df_his["HORAS_HIST"])

# ─── DATOS AGREGADOS ───────────────────────────
horas_hist = df_his.groupby("NOMBRE_INTERINO")["HORAS_HIST"].sum().to_dict()
horas_feb  = df_cob.groupby("NOMBRE_INTERINO")["CARGA_HORARIA"].sum().to_dict()

cols_per_cob = [c for c in ["NOMBRE_INTERINO","CURP","RFC","NSS","TELEFONO","DOMICILIO","COLONIA","CP","MUNICIPIO","CORREO"] if c in df_cob.columns]
cols_per_his = [c for c in ["NOMBRE_INTERINO","CURP","RFC","NSS","TELEFONO","DOMICILIO","COLONIA","CP","MUNICIPIO","CORREO"] if c in df_his.columns]

emp_cob = df_cob[cols_per_cob].drop_duplicates(subset=["NOMBRE_INTERINO"])
emp_his = df_his[cols_per_his].drop_duplicates(subset=["NOMBRE_INTERINO"])
emp_his = emp_his[~emp_his["NOMBRE_INTERINO"].isin(emp_cob["NOMBRE_INTERINO"].tolist())]
empleados_info = pd.concat([emp_cob, emp_his], ignore_index=True).fillna("")

plazas_feb_por_emp = {n: g.to_dict("records") for n, g in df_cob.groupby("NOMBRE_INTERINO")}

# ─── SESSION STATE ─────────────────────────────
if "asignaciones" not in st.session_state:
    st.session_state.asignaciones = {}
if "emp_sel" not in st.session_state:
    st.session_state.emp_sel = None

# ─── HELPERS ───────────────────────────────────
def hrs_nuevas(n): return sum(p["CARGA_HORARIA"] for p in st.session_state.asignaciones.get(n, []))
def hrs_total(n):  return horas_feb.get(n, 0) + hrs_nuevas(n)
def hrs_meta(n):   return horas_hist.get(n, 0)
def hrs_faltan(n): return max(0, hrs_meta(n) - hrs_total(n))

def barra_color(total, meta):
    if meta == 0: return "#94a3b8"
    r = total / meta
    return "#ef4444" if r < 0.5 else ("#f59e0b" if r < 1.0 else "#10b981")

def render_barra(total, meta):
    pct = min(100, total/meta*100) if meta > 0 else 0
    c = barra_color(total, meta)
    return f'<div class="hora-bar-container"><div class="hora-bar" style="width:{pct:.0f}%;background:{c};"></div></div><div style="font-size:0.75rem;color:#64748b;margin-top:0.2rem;">{total} hrs / meta {meta} hrs</div>'

def vacantes_para(nombre):
    usadas = {p["CLAVE_PRESUPUESTAL"] for ps in st.session_state.asignaciones.values() for p in ps}
    actuales = {p["CLAVE_PRESUPUESTAL"] for p in plazas_feb_por_emp.get(nombre, [])}
    excluir = usadas | actuales
    return df_vac[~df_vac["CLAVE_PRESUPUESTAL"].isin(excluir)].to_dict("records")

def generar_oficios():
    oficios = []
    folio = 1

    def agrupar_y_agregar(emp, plazas_lista, es_nueva):
        nonlocal folio
        grupos = {}
        for p in plazas_lista:
            sust = str(p.get("SUSTITUYE_A", p.get("SUSTITUYE A", "")))
            cct  = str(p.get("CLAVE_CCT",   p.get("CLAVE CCT",   "")))
            grupos.setdefault((sust, cct), []).append(p)
        for (titular, cct), grupo in grupos.items():
            horas = sum(int(g.get("CARGA_HORARIA", g.get("CARGA HORARIA", 0))) for g in grupo)
            oficios.append({
                "folio": folio, "emp": emp, "plazas": grupo, "es_nueva": es_nueva,
                "sustituye_a": titular, "clave_cct": cct,
                "nombre_cct":  grupo[0].get("NOMBRE_CCT",   grupo[0].get("NOMBRE CCT",  "")),
                "sostenimiento": grupo[0].get("SOSTENIMIENTO", "ESTATAL"),
                "claves_presupuestales": ", ".join(
                    str(g.get("CLAVE_PRESUPUESTAL", g.get("CLAVE PRESUPUESTAL", ""))) for g in grupo),
                "carga_horaria": f"{horas} HORAS",
                "tipo_alta":    grupo[0].get("TIPO_ALTA",    grupo[0].get("TIPO DE ALTA", "ALTA PROVISIONAL INTERINA")),
                "motivo_vacante": grupo[0].get("MOTIVO_VACANTE", grupo[0].get("MOTIVO VACANTE", "")),
                "plaza":        grupo[0].get("PLAZA", ""),
            })
            folio += 1

    for _, emp in empleados_info.iterrows():
        nombre = emp["NOMBRE_INTERINO"]
        actuales = plazas_feb_por_emp.get(nombre, [])
        if actuales:
            agrupar_y_agregar(emp, actuales, es_nueva=False)
        nuevas = st.session_state.asignaciones.get(nombre, [])
        if nuevas:
            agrupar_y_agregar(emp, nuevas, es_nueva=True)

    return oficios


def fill_plantilla(plantilla_bytes, oficio, fecha_desde, fecha_hasta):
    import zipfile as zf, re as re2
    e = oficio["emp"]
    campos = {
        "FOLIO":                str(oficio["folio"]),
        "NOMBRE_CCT":           str(oficio.get("nombre_cct","")),
        "CLAVE_CCT":            str(oficio.get("clave_cct","")),
        "NOMBRE_INTERINO":      str(e.get("NOMBRE_INTERINO","")),
        "CURP":                 str(e.get("CURP","")),
        "RFC":                  str(e.get("RFC","")),
        "PLAZA":                str(oficio.get("plaza","")),
        "CLAVE_PRESUPUESTAL":   str(oficio.get("claves_presupuestales","")),
        "CARGA_HORARIA":        str(oficio.get("carga_horaria","")),
        "TIPO_DE_ALTA":         str(oficio.get("tipo_alta","")),
        "MOTIVO_VACANTE":       str(oficio.get("motivo_vacante","")),
        "SUSTITUYE_A":          str(oficio.get("sustituye_a","")),
        "NUMERO_SEGURO_SOCIAL": str(e.get("NSS","")),
        "TELÉFONO_MÓVIL":       str(e.get("TELEFONO","")),
        "DOMICILIO":            str(e.get("DOMICILIO","")),
        "COLONIA":              str(e.get("COLONIA","")),
        "CODICO_POSTAL":        str(e.get("CP","")),
        "MUNICIPIO":            str(e.get("MUNICIPIO","")),
        "CORREO_ELECTRÓNICO":   str(e.get("CORREO","")),
    }
    buf = BytesIO()
    with zf.ZipFile(BytesIO(plantilla_bytes), 'r') as zin:
        with zf.ZipFile(buf, 'w', zf.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in ('word/document.xml', 'word/footer1.xml',
                                     'word/footer2.xml', 'word/header1.xml', 'word/header2.xml'):
                    text = data.decode('utf-8')
                    for campo, valor in campos.items():
                        text = text.replace(f'«{campo}»', valor)
                    text = re2.sub(r'\d{1,2} de \w+ de 202\d', fecha_desde, text, count=1)
                    text = re2.sub(r'\d{1,2} de \w+ de 202\d', fecha_hasta, text, count=1)
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    buf.seek(0)
    return buf.read()

def merge_docx_list(docx_bytes_list):
    import zipfile as zf
    from lxml import etree
    if not docx_bytes_list: return b""
    if len(docx_bytes_list) == 1: return docx_bytes_list[0]
    NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def load_parts(b):
        parts = {}
        with zf.ZipFile(BytesIO(b)) as z:
            for n in z.namelist(): parts[n] = z.read(n)
        return parts

    base = load_parts(docx_bytes_list[0])
    base_root = etree.fromstring(base['word/document.xml'])
    base_body = base_root.find(f'{{{NS}}}body')
    children = list(base_body)
    base_sect = children[-1] if children[-1].tag == f'{{{NS}}}sectPr' else None
    if base_sect is not None: base_body.remove(base_sect)

    for docx_bytes in docx_bytes_list[1:]:
        other = load_parts(docx_bytes)
        other_root = etree.fromstring(other['word/document.xml'])
        other_body = other_root.find(f'{{{NS}}}body')
        pb = etree.SubElement(base_body, f'{{{NS}}}p')
        pb_r = etree.SubElement(pb, f'{{{NS}}}r')
        pb_br = etree.SubElement(pb_r, f'{{{NS}}}br')
        pb_br.set(f'{{{NS}}}type', 'page')
        for child in list(other_body):
            if child.tag != f'{{{NS}}}sectPr':
                base_body.append(child)

    if base_sect is not None: base_body.append(base_sect)
    base['word/document.xml'] = etree.tostring(base_root, xml_declaration=True,
                                                encoding='UTF-8', standalone=True)
    buf = BytesIO()
    with zf.ZipFile(buf, 'w', zf.ZIP_DEFLATED) as zout:
        for name, data in base.items(): zout.writestr(name, data)
    buf.seek(0)
    return buf.read()

def generar_word(oficios, plantilla_bytes=None):
    if plantilla_bytes:
        filled = [fill_plantilla(plantilla_bytes, o, fecha_desde, fecha_hasta) for o in oficios]
        return BytesIO(merge_docx_list(filled))
    doc = DocxDocument()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    def p(txt, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
        par = doc.add_paragraph(); par.alignment = align
        par.paragraph_format.space_before = Pt(sb); par.paragraph_format.space_after = Pt(sa)
        r = par.add_run(txt); r.bold = bold; r.font.size = Pt(size)
    def pb():
        par = doc.add_paragraph(); br = OxmlElement('w:br'); br.set(qn('w:type'), 'page')
        par.add_run()._r.append(br)
    for i, o in enumerate(oficios):
        e = o["emp"]
        p("JUAN CARLOS FLORES MIRAMONTES", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p("SECRETARIO DE EDUCACIÓN DEL ESTADO DE JALISCO", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        p("P R E S E N T E", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)
        p("AT´N: NANCY FABIOLA FRANCO ROCHA", size=10)
        p("          DIRECTORA GENERAL DE PERSONAL", size=10, sa=12)
        p(f"El (la) que suscribe Director(a) del centro de trabajo: {o['nombre_cct']} con clave de C.T.: {o['clave_cct']} del sostenimiento {o['sostenimiento']}, me permito presentar a usted la siguiente propuesta en favor de:", sa=10)
        p(f"El (la) C. {e['NOMBRE_INTERINO']}", bold=True)
        p(f"CURP: {e.get('CURP','')}    RFC: {e.get('RFC','')}    Con plaza de: {o['plaza']}", size=10)
        p(f"Clave(s) Presupuestal(es): {o['claves_presupuestales']}", size=10)
        p(f"Sostenimiento: {o['sostenimiento']}", size=10)
        p("Turno: MATUTINO        Carga horaria: " + o['carga_horaria'] + "    Tipo de plaza: HORA/SEMANA/MES", size=10)
        p("Especialidad / Asignatura: NO APLICA", size=10)
        p(f"En calidad de: {o['tipo_alta']}", size=10)
        p(f"Motivo de la vacante: {o['motivo_vacante']}", size=10)
        p(f"Sustituyendo al (el) C. Prof. (a): {o['sustituye_a']}", size=10, sa=12)
        p(f"Con efectos a partir del: {fecha_desde}   Hasta: {fecha_hasta}.", size=10, sa=16)
        p("Sin otro particular de momento, reciba un cordial saludo.", size=10, sa=12)
        p("A T E N T A M E N T E", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=30)
        p("______________________________________________________________", align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        p(nombre_firmante, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
        p(cargo_firmante, align=WD_ALIGN_PARAGRAPH.CENTER)
        if i < len(oficios)-1: pb()
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════
tab1, tab2, tab3 = st.tabs(["📊 Dashboard", "✏️ Asignación de plazas", "📄 Generar propuestas"])

# ─── TAB 1 DASHBOARD ───────────────────────────
with tab1:
    vac_usadas = sum(len(v) for v in st.session_state.asignaciones.values())
    c1,c2,c3,c4 = st.columns(4)
    with c1: st.markdown(f'<div class="metric-card"><div class="value">{len(empleados_info)}</div><div class="label">Empleados</div></div>', unsafe_allow_html=True)
    with c2: st.markdown(f'<div class="metric-card"><div class="value">{len(df_vac)}</div><div class="label">Plazas disponibles</div></div>', unsafe_allow_html=True)
    with c3: st.markdown(f'<div class="metric-card"><div class="value" style="color:#10b981">{sum(1 for _,e in empleados_info.iterrows() if hrs_nuevas(e["NOMBRE_INTERINO"])>0)}</div><div class="label">Con nuevas plazas</div></div>', unsafe_allow_html=True)
    with c4: st.markdown(f'<div class="metric-card"><div class="value" style="color:#f59e0b">{len(df_vac)-vac_usadas}</div><div class="label">Plazas sin asignar</div></div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("Estado por empleado")

    for _, emp in empleados_info.iterrows():
        n = emp["NOMBRE_INTERINO"]
        meta = hrs_meta(n); feb = horas_feb.get(n,0); nuevas = hrs_nuevas(n); total = feb+nuevas; faltan = hrs_faltan(n)
        c1,c2,c3,c4,c5 = st.columns([3,1.3,1.3,2.5,1.2])
        with c1:
            st.markdown(f"**{n}**")
            st.caption(f"RFC: {emp.get('RFC','')} | {emp.get('MUNICIPIO','')}")
        with c2: st.metric("Meta enero", f"{meta} hrs")
        with c3: st.metric("Ya cubre", f"{feb} hrs")
        with c4:
            st.markdown(f"Nuevas: **{nuevas} hrs**")
            st.markdown(render_barra(total, meta), unsafe_allow_html=True)
        with c5:
            if faltan > 0:
                st.markdown(f"<div style='color:#ef4444;font-weight:700'>Faltan {faltan} hrs</div>", unsafe_allow_html=True)
            elif total == meta:
                st.markdown("<div style='color:#10b981;font-weight:700'>✅ Completo</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div style='color:#f59e0b;font-weight:700'>+{total-meta} extra</div>", unsafe_allow_html=True)
            if st.button("Asignar", key=f"b_{n}", use_container_width=True):
                st.session_state.emp_sel = n; st.rerun()
        st.divider()

# ─── TAB 2 ASIGNACIÓN ──────────────────────────
with tab2:
    cl, cd = st.columns([1.3, 2.7])
    with cl:
        st.subheader("Empleados")
        for _, emp in empleados_info.iterrows():
            n = emp["NOMBRE_INTERINO"]
            meta = hrs_meta(n); total = hrs_total(n); faltan = hrs_faltan(n)
            ic = "✅" if faltan==0 and meta>0 else ("🔴" if faltan>0 else "⚪")
            sel = st.session_state.emp_sel == n
            if st.button(f"{ic} {n}\n{total} hrs / meta {meta} hrs", key=f"s_{n}",
                         use_container_width=True, type="primary" if sel else "secondary"):
                st.session_state.emp_sel = n; st.rerun()

    with cd:
        n = st.session_state.emp_sel
        if n is None:
            st.info("👈 Selecciona un empleado para asignarle plazas.")
        else:
            ei = empleados_info[empleados_info["NOMBRE_INTERINO"]==n]
            if ei.empty:
                st.warning("No encontrado.")
            else:
                ei = ei.iloc[0]
                meta=hrs_meta(n); feb=horas_feb.get(n,0); nuevas=hrs_nuevas(n); total=feb+nuevas; faltan=hrs_faltan(n)
                color = barra_color(total, meta)
                pct = min(100, total/meta*100) if meta>0 else 0

                st.markdown(f"### {n}")
                c1,c2,c3 = st.columns(3)
                with c1: st.caption(f"RFC: {ei.get('RFC','')}")
                with c2: st.caption(f"CURP: {ei.get('CURP','')}")
                with c3: st.caption(f"Tel: {ei.get('TELEFONO','')}")

                st.markdown(f"""
                <div style="background:white;border:2px solid {color};border-radius:12px;padding:1rem 1.5rem;margin:0.8rem 0;">
                  <div style="display:flex;gap:2rem;align-items:center;flex-wrap:wrap;">
                    <div><div style="font-size:0.75rem;color:#64748b;font-weight:600;">META ENERO</div>
                         <div style="font-size:2rem;font-weight:800;color:#1a3a5c;">{meta} hrs</div></div>
                    <div><div style="font-size:0.75rem;color:#64748b;font-weight:600;">YA CUBRE</div>
                         <div style="font-size:2rem;font-weight:800;color:#2563a8;">{feb} hrs</div></div>
                    <div><div style="font-size:0.75rem;color:#64748b;font-weight:600;">NUEVAS</div>
                         <div style="font-size:2rem;font-weight:800;color:{color};">{nuevas} hrs</div></div>
                    <div><div style="font-size:0.75rem;color:#64748b;font-weight:600;">FALTAN</div>
                         <div style="font-size:2rem;font-weight:800;color:{'#10b981' if faltan==0 else '#ef4444'};">{'✅' if faltan==0 else str(faltan)+' hrs'}</div></div>
                  </div>
                  <div class="hora-bar-container" style="margin-top:0.8rem;height:16px;">
                    <div class="hora-bar" style="width:{pct:.0f}%;background:{color};"></div>
                  </div>
                </div>""", unsafe_allow_html=True)

                if total > 48:
                    st.error(f"⛔ Total supera 48 hrs ({total} hrs). Quita alguna plaza.")

                p_feb = plazas_feb_por_emp.get(n, [])
                if p_feb:
                    with st.expander(f"📌 Ya cubre en febrero ({len(p_feb)} plazas)"):
                        for p in p_feb:
                            st.markdown(f"""<div class="plaza-item asignada">
                                <div style="font-weight:600;">{p.get('PLAZA','')} — {p.get('CARGA_HORARIA','')} hrs <span class="badge badge-actual">Ya asignada</span></div>
                                <div style="font-size:0.78rem;color:#64748b;">{p.get('CLAVE_PRESUPUESTAL','')}</div>
                                <div style="font-size:0.78rem;color:#64748b;">Sustituye: {p.get('SUSTITUYE_A','')}</div>
                            </div>""", unsafe_allow_html=True)

                st.markdown("#### Plazas nuevas asignadas en esta sesión")
                mis = st.session_state.asignaciones.get(n, [])
                if not mis:
                    st.caption("Ninguna todavía.")
                else:
                    for i, pl in enumerate(mis):
                        cp, cq = st.columns([5,1])
                        with cp:
                            st.markdown(f"""<div class="plaza-item nueva">
                                <div style="font-weight:600;color:#065f46;">{pl.get('PLAZA','')} — {pl.get('CARGA_HORARIA','')} hrs <span class="badge badge-nueva">🆕 Nueva</span></div>
                                <div style="font-size:0.78rem;color:#64748b;">{pl.get('CLAVE_PRESUPUESTAL','')}</div>
                                <div style="font-size:0.78rem;color:#64748b;">Sustituye: {pl.get('SUSTITUYE_A','')}</div>
                                <div style="font-size:0.78rem;color:#64748b;">{pl.get('NOMBRE_CCT','')}</div>
                            </div>""", unsafe_allow_html=True)
                        with cq:
                            if st.button("❌", key=f"rm_{n}_{i}"):
                                st.session_state.asignaciones[n].pop(i); st.rerun()

                disp = vacantes_para(n)
                with st.expander(f"➕ Plazas disponibles para agregar ({len(disp)})"):
                    if not disp:
                        st.caption("No hay plazas disponibles.")
                    else:
                        filtro = st.text_input("Buscar por CCT, titular o clave", key=f"fi_{n}")
                        for pl in disp:
                            clave=str(pl.get("CLAVE_PRESUPUESTAL","")); cct=str(pl.get("NOMBRE_CCT","")); sust=str(pl.get("SUSTITUYE_A",""))
                            if filtro and filtro.upper() not in cct.upper() and filtro.upper() not in sust.upper() and filtro.upper() not in clave.upper():
                                continue
                            hrs_pl = int(pl.get("CARGA_HORARIA",0)); nuevo_total = total+hrs_pl; excede = nuevo_total>48
                            cp, cq = st.columns([5,1])
                            with cp:
                                st.markdown(f"""<div class="plaza-item disponible">
                                    <div style="font-weight:600;">{pl.get('PLAZA','')} — {hrs_pl} hrs</div>
                                    <div style="font-size:0.78rem;color:#64748b;">{clave}</div>
                                    <div style="font-size:0.78rem;color:#64748b;">{cct}</div>
                                    <div style="font-size:0.78rem;color:#64748b;">Sustituye: {sust}</div>
                                    {'<div style="font-size:0.75rem;color:#ef4444;">⚠️ Daría '+str(nuevo_total)+' hrs (máx 48)</div>' if excede else ''}
                                </div>""", unsafe_allow_html=True)
                            with cq:
                                if st.button("➕", key=f"add_{n}_{clave}", disabled=excede):
                                    st.session_state.asignaciones.setdefault(n,[]).append(pl); st.rerun()

# ─── TAB 3 GENERAR ─────────────────────────────
with tab3:
    oficios = generar_oficios()
    if not oficios:
        st.warning("No hay plazas asignadas aún.")
    else:
        st.success(f"✅ Se generarán **{len(oficios)} oficio(s)**.")
        relacion = [{"Oficio #": o["folio"],
                     "Tipo": "🆕 Nueva" if o.get("es_nueva") else "📌 Actual",
                     "Empleado": o["emp"]["NOMBRE_INTERINO"],
                     "RFC": o["emp"].get("RFC",""), "Carga Horaria": o["carga_horaria"],
                     "Sustituye a": o["sustituye_a"], "CCT": o["clave_cct"],
                     "Vigencia": f"{fecha_desde} – {fecha_hasta}"} for o in oficios]
        st.subheader("Relación para RH Oficinas Centrales")
        st.dataframe(pd.DataFrame(relacion), use_container_width=True, hide_index=True)

        st.subheader("Vista previa de oficios")
        for o in oficios:
            with st.expander(f"Oficio #{o['folio']} — {o['emp']['NOMBRE_INTERINO']} — {o['carga_horaria']}"):
                st.markdown(f"""<div class="oficio-preview">
                    <b>Empleado:</b> {o['emp']['NOMBRE_INTERINO']}<br>
                    <b>CURP:</b> {o['emp'].get('CURP','')} &nbsp; <b>RFC:</b> {o['emp'].get('RFC','')}<br>
                    <b>Plaza:</b> {o['plaza']}<br>
                    <b>Clave(s) Presupuestal(es):</b> {o['claves_presupuestales']}<br>
                    <b>Carga horaria:</b> {o['carga_horaria']}<br>
                    <b>CCT:</b> {o['nombre_cct']} ({o['clave_cct']})<br>
                    <b>Motivo vacante:</b> {o['motivo_vacante']}<br>
                    <b>Sustituye a:</b> {o['sustituye_a']}<br>
                    <b>Tipo de alta:</b> {o['tipo_alta']}<br>
                    <b>Vigencia:</b> {fecha_desde} al {fecha_hasta}
                </div>""", unsafe_allow_html=True)

        st.divider()
        cw, cx, cc = st.columns(3)
        with cw:
            if st.button("📄 Generar Word", type="primary", use_container_width=True):
                with st.spinner("Generando..."):
                    plt_bytes = f_plantilla.read() if f_plantilla else None
                    buf = generar_word(oficios, plt_bytes)
                st.download_button("⬇️ Descargar propuestas.docx", data=buf,
                    file_name="propuestas_plazas.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
        with cx:
            if st.button("📊 Generar Excel de propuestas", type="primary", use_container_width=True):
                filas = []
                folio_excel = 1
                for o in oficios:
                    e = o["emp"]
                    for pl in o["plazas"]:
                        filas.append({
                            "ORIGEN": "NUEVA ASIGNACIÓN" if o.get("es_nueva") else "CONTINUIDAD (COBERTURA)",
                            "FOLIO": folio_excel,
                            "SOSTENIMIENTO": pl.get("SOSTENIMIENTO","ESTATAL"),
                            "CLAVE CCT": pl.get("CLAVE_CCT",""),
                            "NOMBRE CCT": pl.get("NOMBRE_CCT",""),
                            "EFECTO INICIAL": fecha_desde,
                            "EFECTO FINAL": fecha_hasta,
                            "PLAZA": pl.get("PLAZA",""),
                            "CLAVE PRESUPUESTAL": pl.get("CLAVE_PRESUPUESTAL",""),
                            "TURNO": pl.get("TURNO","MATUTINO"),
                            "CARGA HORARIA": pl.get("CARGA_HORARIA",""),
                            "TIPO DE PLAZA": pl.get("TIPO DE PLAZA", pl.get("TIPO_PLAZA","HORA/SEMANA/MES")),
                            "TIPO DE ALTA": pl.get("TIPO_ALTA","ALTA PROVISIONAL INTERINA"),
                            "MOTIVO VACANTE": pl.get("MOTIVO_VACANTE",""),
                            "SUSTITUYE A": pl.get("SUSTITUYE_A",""),
                            "NOMBRE INTERINO": e.get("NOMBRE_INTERINO",""),
                            "CURP": e.get("CURP",""),
                            "RFC": e.get("RFC",""),
                            "NUMERO SEGURO SOCIAL": e.get("NSS",""),
                            "TELÉFONO MÓVIL": e.get("TELEFONO",""),
                            "DOMICILIO": e.get("DOMICILIO",""),
                            "COLONIA": e.get("COLONIA",""),
                            "CODICO POSTAL": e.get("CP",""),
                            "MUNICIPIO": e.get("MUNICIPIO",""),
                            "CORREO ELECTRÓNICO": e.get("CORREO",""),
                        })
                        folio_excel += 1
                df_out = pd.DataFrame(filas)
                xbuf = BytesIO()
                with pd.ExcelWriter(xbuf, engine="openpyxl") as writer:
                    df_out.to_excel(writer, index=False, sheet_name="Propuestas")
                    ws = writer.sheets["Propuestas"]
                    from openpyxl.styles import PatternFill, Font, Alignment
                    fill = PatternFill("solid", fgColor="1A3A5C")
                    font = Font(color="FFFFFF", bold=True, size=10)
                    for cell in ws[1]:
                        cell.fill = fill; cell.font = font
                        cell.alignment = Alignment(horizontal="center", wrap_text=True)
                    for col in ws.columns:
                        max_len = max(len(str(c.value or "")) for c in col)
                        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 35)
                xbuf.seek(0)
                st.download_button("⬇️ Descargar propuestas.xlsx", data=xbuf,
                    file_name="propuestas_plazas.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
        with cc:
            cb = BytesIO()
            pd.DataFrame(relacion).to_csv(cb, index=False, encoding="utf-8-sig"); cb.seek(0)
            st.download_button("📋 Descargar relación RH (CSV)", data=cb,
                file_name="relacion_rh.csv", mime="text/csv", use_container_width=True)

        st.divider()
        if st.button("💾 Guardar asignaciones nuevas en Google Sheets", use_container_width=True):
            nuevas = [o for o in oficios if o.get("es_nueva")]
            if not nuevas:
                st.warning("No hay plazas nuevas para guardar.")
            else:
                with st.spinner("Guardando en Google Sheets..."):
                    ok = guardar_asignaciones(oficios, fecha_desde, fecha_hasta)
                if ok:
                    st.success(f"✅ {len(nuevas)} oficio(s) nuevo(s) guardados en la hoja ASIGNACIONES.")